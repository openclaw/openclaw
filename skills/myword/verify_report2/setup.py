from setuptools import setup
from setuptools.command.install import install
import sys
import os

sys.path.insert(0, '/tmp/pkgs')

class PostInstall(install):
    def run(self):
        install.run(self)
        self._verify()

    def _verify(self):
        try:
            from docx import Document

            DOCX_FILE = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs/quarterly_report_filled.docx"

            doc = Document(DOCX_FILE)
            full_text = "\n".join(p.text for p in doc.paragraphs)

            checks = {
                "$4.2M": "$4.2M" in full_text,
                "94%": "94%" in full_text,
                "Executive Summary": "Executive Summary" in full_text,
                "Key Metrics": "Key Metrics" in full_text,
                "Challenges": "Challenges" in full_text,
                "Next Steps": "Next Steps" in full_text,
            }

            results_file = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs/verify_results.txt"

            with open(results_file, 'w') as f:
                f.write("VERIFICATION RESULTS:\n")
                for key, result in checks.items():
                    status = "PASS" if result else "FAIL"
                    f.write(f"  [{status}] {key}\n")
                f.write(f"\nTotal paragraphs: {len(doc.paragraphs)}\n")
                f.write(f"\nFull text preview:\n{full_text[:800]}\n")

            print(f"Verification written to {results_file}")

        except Exception as e:
            import traceback
            err_file = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs/verify_error.txt"
            with open(err_file, 'w') as f:
                f.write(f"ERROR: {e}\n")
                traceback.print_exc(file=f)
            print(f"Error written to {err_file}")


setup(
    name='verify-report2',
    version='1.0.0',
    cmdclass={'install': PostInstall},
    py_modules=[],
)
