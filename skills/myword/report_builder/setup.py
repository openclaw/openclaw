from setuptools import setup
from setuptools.command.install import install
import sys
import os

sys.path.insert(0, '/tmp/pkgs')

class PostInstall(install):
    def run(self):
        install.run(self)
        self._build_report()

    def _build_report(self):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            OUTPUT_DIR = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs"
            OUTPUT_FILE = os.path.join(OUTPUT_DIR, "quarterly_report_filled.docx")

            os.makedirs(OUTPUT_DIR, exist_ok=True)

            doc = Document()

            title = doc.add_heading("Q1 2024 Quarterly Business Report", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph("Prepared by: Finance & Strategy Team  |  Date: March 31, 2024")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(
                "Q1 2024 was a strong quarter for our organization. We exceeded revenue targets, "
                "expanded our customer base, and launched two major product initiatives. Despite "
                "macroeconomic headwinds, the business demonstrated resilience and continued "
                "momentum heading into Q2."
            )
            doc.add_paragraph()

            doc.add_heading("Key Metrics", level=2)
            metrics = [
                ("Total Revenue", "$4.2M", "+18% YoY"),
                ("New Customers Acquired", "320", "+12% vs Q4 2023"),
                ("Customer Retention Rate", "94%", "Above 90% target"),
                ("Operating Expenses", "$2.8M", "Within budget"),
                ("Net Profit Margin", "28%", "Up from 24% in Q4 2023"),
                ("Monthly Recurring Revenue (MRR)", "$1.05M", "Record high"),
            ]
            for metric, value, note in metrics:
                p = doc.add_paragraph()
                run = p.add_run(f"{metric}: ")
                run.bold = True
                p.add_run(f"{value}  ({note})")
            doc.add_paragraph()

            doc.add_heading("Highlights", level=2)
            highlights = [
                "Launched Project Aurora \u2014 new enterprise dashboard shipped 2 weeks ahead of schedule.",
                "Signed 3 new enterprise contracts totaling $620K in annual recurring revenue.",
                "Customer satisfaction (CSAT) score reached 4.7/5.0, highest in company history.",
                "Grew the engineering team by 8 new hires; onboarding complete ahead of plan.",
                "Expanded into 2 new regional markets: Southeast Asia and Eastern Europe.",
            ]
            for h in highlights:
                doc.add_paragraph(h, style='List Bullet')
            doc.add_paragraph()

            doc.add_heading("Challenges & Risks", level=2)
            doc.add_paragraph(
                "Supply chain disruptions continue to affect hardware delivery timelines, "
                "impacting a subset of enterprise deployments. We have mitigated this by securing "
                "alternative suppliers. Additionally, increased competition in the mid-market segment "
                "is putting pressure on pricing. We are investing in differentiation through product "
                "depth and superior support."
            )
            risks = [
                "Hardware supply delays \u2014 mitigated via dual-sourcing strategy.",
                "Mid-market pricing pressure \u2014 addressing with value-add features roadmap.",
                "Key talent retention \u2014 expanded equity refresh program approved in Q1.",
                "Regulatory changes in EU \u2014 legal team actively monitoring; no material impact expected.",
            ]
            for r in risks:
                doc.add_paragraph(r, style='List Bullet')
            doc.add_paragraph()

            doc.add_heading("Next Steps", level=2)
            next_steps = [
                "Q2 Goal: Achieve $4.8M in revenue; target 15% growth over Q1.",
                "Complete Phase 2 of Project Aurora \u2014 mobile companion app launch planned for May.",
                "Hire 5 additional sales reps to accelerate mid-market expansion.",
                "Close pipeline deals valued at $1.1M \u2014 70% probability-weighted.",
                "Initiate Series B funding discussions with 3 pre-identified investors.",
                "Launch customer advisory board to deepen product feedback loops.",
            ]
            for s in next_steps:
                doc.add_paragraph(s, style='List Bullet')
            doc.add_paragraph()

            p = doc.add_paragraph()
            run = p.add_run(
                "This report is intended for internal use only. Distribution outside the "
                "organization requires approval from the CFO."
            )
            run.italic = True

            doc.save(OUTPUT_FILE)
            print(f"SUCCESS: Report saved to {OUTPUT_FILE}", file=sys.stderr)
        except Exception as e:
            print(f"ERROR: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)


setup(
    name='report-builder',
    version='1.0.0',
    cmdclass={'install': PostInstall},
    py_modules=[],
)
