from setuptools import setup
from setuptools.command.install import install
import sys
import os

sys.path.insert(0, '/tmp/pkgs')

OUTPUTS = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs"

class PostInstall(install):
    def run(self):
        install.run(self)
        self._verify()

    def _verify(self):
        results = []
        try:
            from docx import Document

            DOCX_FILE = os.path.join(OUTPUTS, "quarterly_report_filled.docx")
            doc = Document(DOCX_FILE)
            full_text = "\n".join(p.text for p in doc.paragraphs)

            checks = {
                "$4.2M": "$4.2M" in full_text,
                "94%": "94%" in full_text,
                "Executive Summary": "Executive Summary" in full_text,
                "Key Metrics": "Key Metrics" in full_text,
                "Challenges": "Challenges" in full_text,
                "Next Steps": "Next Steps" in full_text,
            }

            results.append("VERIFICATION RESULTS:")
            for key, result in checks.items():
                status = "PASS" if result else "FAIL"
                results.append(f"  [{status}] {key}")
            results.append(f"\nTotal paragraphs: {len(doc.paragraphs)}")
            results.append(f"\nFull text preview:\n{full_text[:1000]}")

        except Exception as e:
            import traceback
            results.append(f"ERROR: {e}")
            results.append(traceback.format_exc())

        results_file = os.path.join(OUTPUTS, "verify_results.txt")
        with open(results_file, 'w') as f:
            f.write("\n".join(results))


setup(
    name='verify-report3',
    version='1.0.0',
    cmdclass={'install': PostInstall},
    py_modules=[],
)
