from setuptools import setup
from setuptools.command.install import install
import sys
import os

sys.path.insert(0, '/tmp/pkgs')

class PostInstall(install):
    def run(self):
        install.run(self)
        self._verify()

    def _verify(self):
        try:
            from docx import Document

            DOCX_FILE = "/Users/denizburcayhaberal/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/07087d0e-14d6-4635-ac59-cb7a1426a26a/0180c3cf-14ac-4364-b704-5c6dc6dbc8e8/skills/openclaw-template-filler-workspace/iteration-1/eval-2-english-report-with-file/without_skill/outputs/quarterly_report_filled.docx"

            doc = Document(DOCX_FILE)
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"

            checks = {
                "$4.2M": "$4.2M" in full_text,
                "94%": "94%" in full_text,
                "Executive Summary": "Executive Summary" in full_text,
                "Key Metrics": "Key Metrics" in full_text,
                "Challenges & Risks": "Challenges" in full_text,
                "Next Steps": "Next Steps" in full_text,
            }

            print("VERIFICATION RESULTS:", file=sys.stderr)
            for key, result in checks.items():
                status = "PASS" if result else "FAIL"
                print(f"  [{status}] {key}", file=sys.stderr)

            print(f"\nTotal paragraphs: {len(doc.paragraphs)}", file=sys.stderr)
            print(f"Full text preview (first 500 chars):", file=sys.stderr)
            print(full_text[:500], file=sys.stderr)

        except Exception as e:
            print(f"VERIFY ERROR: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)


setup(
    name='verify-report',
    version='1.0.0',
    cmdclass={'install': PostInstall},
    py_modules=[],
)
