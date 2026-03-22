"""
标点恢复 + Word 格式化脚本
用法：python format_docx.py "<原始无标点txt路径>" "<输出docx路径>"
依赖：pip install python-docx
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TXT_PATH  = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\Administrator\.openclaw\workspace\raw_text.txt"
OUT_PATH  = sys.argv[2] if len(sys.argv) > 2 else r"C:\Users\Administrator\Desktop\video_transcript_final.docx"
TITLE     = sys.argv[3] if len(sys.argv) > 3 else "视频转录文字整理"

with open(TXT_PATH, "r", encoding="utf-8") as f:
    text = f.read().strip()

doc = Document()

# 标题
title = doc.add_heading(TITLE, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题（可选）
sub = doc.add_paragraph("——视频转录文字整理")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 正文：按空行分段
for para in text.split("\n\n"):
    if not para.strip():
        continue
    p = doc.add_paragraph(para.strip())
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(24)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
print(f"Characters: {len(text)}, Paragraphs: {len(text.split(chr(10)+chr(10)))}")
