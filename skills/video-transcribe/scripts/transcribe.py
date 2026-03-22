"""
视频转文字脚本（Whisper medium 模型）
用法：python transcribe.py "<视频路径>" "<输出docx路径>"
依赖：pip install openai-whisper python-docx
"""
import sys
import whisper
from docx import Document

VIDEO_PATH = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\Administrator\Desktop\1.mp4"
OUT_PATH   = sys.argv[2] if len(sys.argv) > 2 else r"C:\Users\Administrator\Desktop\video_transcript_raw.docx"

print(f"Loading model (medium)...", flush=True)
model = whisper.load_model("medium")
print("Model loaded", flush=True)

print(f"Transcribing: {VIDEO_PATH}", flush=True)
result = model.transcribe(VIDEO_PATH, language="zh")
text = result["text"]
print(f"Done. Characters: {len(text)}", flush=True)

doc = Document()
doc.add_heading("视频转写文字", 0)
doc.add_paragraph(text)
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}", flush=True)
